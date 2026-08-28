import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

doc_path = r'C:\Users\Yoei\Documents\2026.8實習資料\傑海達研究紀錄簿_2026-08_林千佑.docx'
doc = docx.Document(doc_path)
t9 = doc.tables[9]
cell = t9.rows[1].cells[0]

updates = [
    {
        'p_idx': 9,
        'rId': 'rId57',
        'link_text': '[75046ffd]',
        'url': 'https://github.com/hereyougoai/moos-ivp/commit/75046ffdc3e4648d1b7eb581d36f9b82c7c29199',
        'suffix_text': ' fix(target): raise default target max_speed ceiling from 1.2 to 3.0 m/s and add max_spd argument'
    },
    {
        'p_idx': 10,
        'rId': 'rId58',
        'link_text': '[176ab4f5]',
        'url': 'https://github.com/hereyougoai/moos-ivp/commit/176ab4f58b405d58c186ffeb7be8a5310d1a7546',
        'suffix_text': ' feat(target): implement 8 target scenario behavior profiles and shoreside live profile switching menu'
    },
    {
        'p_idx': 11,
        'rId': 'rId59',
        'link_text': '[b1209e11]',
        'url': 'https://github.com/hereyougoai/moos-ivp/commit/b1209e114fa3ec7954755f5cee3df0c81bffca2e',
        'suffix_text': ' feat(pTargetCoordinator): add target course smoothing, lead hysteresis, latched retreat station, and rate limiting'
    },
    {
        'p_idx': 12,
        'rId': 'rId60',
        'link_text': '[cdc77dfe]',
        'url': 'https://github.com/hereyougoai/moos-ivp/commit/cdc77dfe86c389e71e16bb892d82b56306d89f2d',
        'suffix_text': ' test(m_shield_demo): add profile sweep benchmark script and escape path analyzer'
    },
    {
        'p_idx': 13,
        'rId': 'rId61',
        'link_text': '[6c819811]',
        'url': 'https://github.com/hereyougoai/moos-ivp/commit/6c819811324b3b9374b81d1cea23bdc7188e46fa',
        'suffix_text': ' docs: update mission documentation and dev_log_20260827 with target profiles and speed ceiling specs'
    },
    {
        'p_idx': 15,
        'rId': 'rId62',
        'link_text': '▶ 點此前往 Google Drive 觀看【115/8/27 當日錄影資料夾】',
        'url': 'https://drive.google.com/drive/u/0/folders/1SHnjHqxBRG87igXT-ngJhdldMeZj2ItF',
        'suffix_text': None
    }
]

for up in updates:
    p = cell.paragraphs[up['p_idx']]
    rId = up['rId']
    if rId in doc.part.rels:
        doc.part.rels[rId]._target = up['url']
        print(f"Updated {rId} target to {up['url']}")
    
    for child in p._element:
        if child.tag.endswith('hyperlink') and child.get(qn('r:id')) == rId:
            child.clear()
            child.set(qn('r:id'), rId)
            color_val = 'D9381E' if up['p_idx'] == 15 else '0066CC'
            r_elem = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:b/><w:color w:val="{color_val}"/><w:u w:val="single"/></w:rPr><w:t>{up["link_text"]}</w:t></w:r>')
            child.append(r_elem)
        elif child.tag.endswith('r') and up['suffix_text'] is not None:
            t_elem = child.find(qn('w:t'))
            if t_elem is not None and any(k in t_elem.text for k in ['feat', 'fix', 'test', 'docs']):
                t_elem.text = up['suffix_text']

doc.save(doc_path)
print("Successfully saved docx to:", doc_path)
